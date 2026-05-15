"""二进制文件解析器：PDF / Word。

注意：底层第三方库（pdfplumber、python-docx）可能包含已知安全漏洞（CVE）。
这些库直接解析二进制格式，在不受信任的输入上使用时存在风险。
如需处理不受信任的文件，建议在隔离环境中运行。
"""

import logging
from pathlib import Path
from .base import FileParser, ParsedDocument

logger = logging.getLogger(__name__)


class BinaryFileParser(FileParser):
    supported_extensions = [".pdf", ".docx"]

    def parse(self, path: Path) -> ParsedDocument:
        ext = path.suffix.lower()
        if ext == ".pdf":
            return self._parse_pdf(path)
        else:
            return self._parse_docx(path)

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("PDF 解析需要 pdfplumber，请执行: pip install pdfplumber")
        sections = []
        raw_parts = []
        try:
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        sections.append({"heading": f"Page {i+1}", "content": text})
                        raw_parts.append(text)
        except Exception as e:
            logger.error("PDF 解析失败 (%s): %s", path, e, exc_info=True)
            raise RuntimeError(
                f"PDF 文件解析失败: {path.name}。文件可能已损坏或使用了不支持的编码格式。"
                f" 底层错误: {e}"
            ) from e
        return ParsedDocument(
            source_path=path, format="pdf", title=path.stem,
            sections=sections, raw_text="\n".join(raw_parts),
            metadata={"pages": len(sections)},
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Word 解析需要 python-docx，请执行: pip install python-docx")
        try:
            doc = Document(str(path))
            raw_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    raw_parts.append(para.text)
            raw_text = "\n".join(raw_parts)
        except Exception as e:
            logger.error("DOCX 解析失败 (%s): %s", path, e, exc_info=True)
            raise RuntimeError(
                f"Word 文档解析失败: {path.name}。文件可能已损坏或使用了不支持的格式。"
                f" 底层错误: {e}"
            ) from e
        return ParsedDocument(
            source_path=path, format="word", title=path.stem,
            sections=[{"heading": path.stem, "content": raw_text}],
            raw_text=raw_text,
        )
